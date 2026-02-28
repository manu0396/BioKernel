import sys
import os
import re
from pathlib import Path

try:
    from docx import Document
except ImportError:
    raise SystemExit("Falta python-docx. Instala con: python -m pip install python-docx")

NOTE = (
    "NOTA (actualizado 2026-02-26): La actividad se desarrolla actualmente como AUTÓNOMO "
    "(fase pre-S.L.U.), con base en Manzanares el Real (Madrid, España) y prioridad 100% REMOTO. "
    "La constitución de una S.L.U. se realizará únicamente cuando el proyecto sea realmente rentable "
    "(p. ej., ingresos recurrentes, margen suficiente y/o necesidad de contratación/financiación)."
)

# Reemplazos “seguros” (no agresivos)
REPLACEMENTS = [
    # Variantes SLU
    (re.compile(r"\bNeogenesis\s*S\.L\.U\.?\b", re.IGNORECASE), "NeoGenesis (autónomo, fase pre-S.L.U.)"),
    (re.compile(r"\bRegenOps\s*S\.L\.U\.?\b", re.IGNORECASE), "RegenOps (autónomo, fase pre-S.L.U.)"),
    (re.compile(r"\bS\.L\.U\.?\b"), "S.L.U. (prevista cuando sea rentable)"),
    (re.compile(r"\bSLU\b"), "S.L.U. (prevista cuando sea rentable)"),

    # Placeholders típicos
    (re.compile(r"\[DIRECCI[ÓO]N\]|\[DIRECCION\]", re.IGNORECASE), "Manzanares el Real (Madrid, España)"),
    (re.compile(r"\[LOCALIDAD\]|\[CIUDAD\]|\[MUNICIPIO\]", re.IGNORECASE), "Manzanares el Real"),
]

def apply_replacements_text(text: str) -> str:
    out = text
    for pattern, repl in REPLACEMENTS:
        out = pattern.sub(repl, out)
    return out

def docx_replace_in_paragraph(paragraph) -> int:
    """Devuelve nº de cambios aproximado."""
    original = paragraph.text
    if not original:
        return 0

    updated = apply_replacements_text(original)
    if updated == original:
        return 0

    # Intentar reemplazo run a run (mantiene estilos cuando no hay cortes raros)
    changed = 0
    for run in paragraph.runs:
        if run.text:
            new_run = apply_replacements_text(run.text)
            if new_run != run.text:
                run.text = new_run
                changed += 1

    # Si el texto total sigue sin coincidir (patrón cruzando runs), colapsar el párrafo
    if paragraph.text != updated:
        # borrar runs y poner texto plano
        for run in paragraph.runs:
            run.text = ""
        paragraph.add_run(updated)
        changed += 1

    return changed

def docx_replace_in_table(table) -> int:
    c = 0
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                c += docx_replace_in_paragraph(p)
            for t in cell.tables:
                c += docx_replace_in_table(t)
    return c

def ensure_note_in_docx(doc: "Document") -> bool:
    # Evitar duplicar la nota
    joined = "\n".join([p.text for p in doc.paragraphs[:8] if p.text])
    if "fase pre" in joined.lower() or "manzanares el real" in joined.lower():
        return False

    # Insertar al inicio (más robusto)
    if doc.paragraphs:
        doc.paragraphs[0].insert_paragraph_before(NOTE)
    else:
        doc.add_paragraph(NOTE)
    return True

def process_docx(path: Path) -> dict:
    doc = Document(str(path))
    inserted = ensure_note_in_docx(doc)

    changes = 0
    for p in doc.paragraphs:
        changes += docx_replace_in_paragraph(p)
    for t in doc.tables:
        changes += docx_replace_in_table(t)

    # Headers/footers (si los hay)
    for section in doc.sections:
        for hf in [section.header, section.footer]:
            for p in hf.paragraphs:
                changes += docx_replace_in_paragraph(p)
            for t in hf.tables:
                changes += docx_replace_in_table(t)

    if inserted or changes > 0:
        doc.save(str(path))
    return {"file": str(path), "inserted_note": inserted, "changes": changes}

def ensure_note_in_text_file(content: str) -> str:
    if "fase pre" in content.lower() or "manzanares el real" in content.lower():
        return content
    return NOTE + "\n\n" + content

def process_text_file(path: Path) -> dict:
    try:
        content = path.read_text(encoding="utf-8", errors="ignore")
    except Exception:
        return {"file": str(path), "updated": False, "changes": 0}

    original = content
    content = apply_replacements_text(content)
    content = ensure_note_in_text_file(content)

    if content != original:
        path.write_text(content, encoding="utf-8")
        return {"file": str(path), "updated": True, "changes": 1}
    return {"file": str(path), "updated": False, "changes": 0}

def main():
    if len(sys.argv) < 2:
        print("Uso: python fix_docx_autonomo_preslu.py <DIRECTORIO>")
        raise SystemExit(2)

    root = Path(sys.argv[1]).resolve()
    if not root.exists() or not root.is_dir():
        raise SystemExit(f"Directorio inválido: {root}")

    report = []
    docx_count = 0
    text_count = 0

    for p in root.rglob("*"):
        if not p.is_file():
            continue
        ext = p.suffix.lower()
        if ext == ".docx":
            docx_count += 1
            report.append(process_docx(p))
        elif ext in (".md", ".txt", ".csv"):
            text_count += 1
            report.append(process_text_file(p))

    # Reporte
    report_path = root / "__AUTOFIX_REPORT.md"
    updated = [r for r in report if (r.get("changes", 0) or r.get("inserted_note") or r.get("updated"))]
    lines = []
    lines.append("# Reporte de actualización (autónomo pre-S.L.U.)")
    lines.append("")
    lines.append(f"- Directorio: `{root}`")
    lines.append(f"- DOCX procesados: {docx_count}")
    lines.append(f"- Textos procesados (md/txt/csv): {text_count}")
    lines.append(f"- Archivos modificados: {len(updated)}")
    lines.append("")
    lines.append("## Archivos modificados")
    for r in updated:
        lines.append(f"- {r['file']}")
    report_path.write_text("\n".join(lines), encoding="utf-8")

    print(f"OK. Reporte: {report_path}")

if __name__ == "__main__":
    main()